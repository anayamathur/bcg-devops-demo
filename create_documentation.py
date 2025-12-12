#!/usr/bin/env python3
"""
BCG DevOps GenAI POC - Documentation Generator
Creates a comprehensive Word document with architecture diagrams and workflow explanations.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

def create_documentation():
    # Create document
    doc = Document()
    
    # Set document title
    title = doc.add_heading('BCG DevOps GenAI POC', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Autonomous DevOps Agent with AI-Powered Code Analysis')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Version and date info
    info = doc.add_paragraph()
    info.add_run('Version: ').bold = True
    info.add_run('1.0\n')
    info.add_run('Date: ').bold = True
    info.add_run('December 2024\n')
    info.add_run('Project: ').bold = True
    info.add_run('BCG DevOps GenAI Proof of Concept')
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # Table of Contents
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Executive Summary',
        '2. System Architecture Overview',
        '3. Core Components',
        '4. API Endpoints',
        '5. Autonomous DevOps Agent Workflow',
        '6. CI/CD Pipeline Integration',
        '7. Security Features',
        '8. Technology Stack',
        '9. Deployment Guide',
        '10. Usage Examples'
    ]
    for item in toc_items:
        doc.add_paragraph(item)
    
    doc.add_page_break()
    
    # 1. Executive Summary
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        'The BCG DevOps GenAI POC is an innovative proof-of-concept that demonstrates the power of '
        'integrating Generative AI into DevOps workflows. This solution leverages Amazon Bedrock\'s '
        'Nova Pro model to create an autonomous DevOps agent capable of:'
    )
    
    features = [
        'Analyzing code repositories for patterns, issues, and improvements',
        'Generating production-ready code based on natural language requirements',
        'Performing comprehensive security scans using Trivy and Bandit',
        'Automatically fixing identified vulnerabilities with AI-generated patches',
        'Creating pull requests with detailed descriptions and change summaries',
        'Providing an interactive chat interface for DevOps assistance',
        'Running complete DevOps pipelines autonomously'
    ]
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph(
        'This solution significantly reduces manual intervention in DevOps workflows, accelerates '
        'development cycles, and ensures consistent security practices across all code changes.'
    )
    
    # 2. System Architecture Overview
    doc.add_page_break()
    doc.add_heading('2. System Architecture Overview', level=1)
    
    doc.add_paragraph(
        'The system follows a serverless architecture pattern, leveraging AWS services for '
        'scalability, reliability, and cost-efficiency.'
    )
    
    # Add architecture diagram
    diagram_path = 'generated-diagrams/bcg-system-architecture.png'
    if os.path.exists(diagram_path):
        doc.add_paragraph()
        doc.add_picture(diagram_path, width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption = doc.add_paragraph('Figure 1: System Architecture Overview')
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.runs[0].italic = True
    
    doc.add_paragraph()
    doc.add_heading('Architecture Components:', level=2)
    
    components = [
        ('Frontend (Web Dashboard)', 'A responsive HTML/JavaScript single-page application that provides an intuitive interface for interacting with the DevOps agent.'),
        ('API Gateway', 'AWS API Gateway serves as the entry point, handling REST API requests, authentication, and rate limiting.'),
        ('Lambda Function', 'A Python-based serverless function (5,000+ lines) that processes all requests and orchestrates the DevOps workflows.'),
        ('Amazon Bedrock', 'Provides the AI backbone using the Nova Pro model for code analysis, generation, and natural language understanding.'),
        ('GitHub Integration', 'Direct API integration for repository operations, PR creation, and code management.'),
        ('Security Tools', 'Integration with Trivy and Bandit for comprehensive security scanning.')
    ]
    
    for name, desc in components:
        p = doc.add_paragraph()
        p.add_run(f'{name}: ').bold = True
        p.add_run(desc)
    
    # 3. Core Components
    doc.add_page_break()
    doc.add_heading('3. Core Components', level=1)
    
    doc.add_heading('3.1 Lambda Function (index.py)', level=2)
    doc.add_paragraph(
        'The Lambda function is the heart of the system, containing over 5,000 lines of Python code '
        'that handles all DevOps operations. Key modules include:'
    )
    
    modules = [
        ('Request Handler', 'Routes incoming API requests to appropriate handlers'),
        ('Repository Analyzer', 'Clones and analyzes GitHub repositories'),
        ('Code Generator', 'Uses AI to generate code based on requirements'),
        ('Security Scanner', 'Integrates with Trivy and Bandit for vulnerability detection'),
        ('Auto-Fixer', 'Generates AI-powered fixes for identified issues'),
        ('PR Creator', 'Creates well-documented pull requests on GitHub'),
        ('Chat Handler', 'Provides conversational AI interface for DevOps queries'),
        ('Autonomous Agent', 'Orchestrates complete DevOps workflows automatically')
    ]
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    header = table.rows[0].cells
    header[0].text = 'Module'
    header[1].text = 'Description'
    for cell in header:
        cell.paragraphs[0].runs[0].bold = True
    
    for module, desc in modules:
        row = table.add_row().cells
        row[0].text = module
        row[1].text = desc
    
    doc.add_paragraph()
    
    doc.add_heading('3.2 AI Integration (Amazon Bedrock)', level=2)
    doc.add_paragraph(
        'The system uses Amazon Bedrock\'s Nova Pro model (amazon.nova-pro-v1:0) for all AI operations. '
        'Key AI capabilities include:'
    )
    
    ai_features = [
        'Natural language understanding for parsing user requirements',
        'Code analysis and pattern recognition',
        'Vulnerability assessment and prioritization',
        'Code generation following best practices',
        'Automated fix generation for security issues',
        'Contextual conversation for DevOps assistance'
    ]
    for feature in ai_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_heading('3.3 Frontend Dashboard', level=2)
    doc.add_paragraph(
        'The frontend is a single HTML file with embedded CSS and JavaScript, providing:'
    )
    
    frontend_features = [
        'Repository analysis interface with real-time results',
        'Code generation with syntax highlighting',
        'Security scan visualization with severity indicators',
        'Interactive chat interface for AI assistance',
        'Pipeline execution controls and monitoring',
        'Pull request creation and tracking'
    ]
    for feature in frontend_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    # 4. API Endpoints
    doc.add_page_break()
    doc.add_heading('4. API Endpoints', level=1)
    
    # Add API flow diagram
    diagram_path = 'generated-diagrams/api-request-flow.png'
    if os.path.exists(diagram_path):
        doc.add_picture(diagram_path, width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption = doc.add_paragraph('Figure 2: API Request Flow')
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.runs[0].italic = True
    
    doc.add_paragraph()
    doc.add_paragraph('Base URL: https://4dyb4z9kgk.execute-api.us-east-1.amazonaws.com/prod')
    doc.add_paragraph()
    
    endpoints = [
        ('GET /health', 'Health check endpoint', 'Returns system status and version'),
        ('POST /analyze', 'Repository analysis', 'Analyzes a GitHub repository for structure, languages, and patterns'),
        ('POST /generate', 'Code generation', 'Generates code based on natural language requirements'),
        ('POST /scan', 'Security scanning', 'Performs security scan using Trivy and Bandit'),
        ('POST /fix', 'Auto-fix vulnerabilities', 'Generates AI-powered fixes for security issues'),
        ('POST /chat', 'AI chat interface', 'Interactive conversation for DevOps assistance'),
        ('POST /autonomous', 'Autonomous agent', 'Runs complete DevOps workflow automatically'),
        ('POST /pipeline', 'Pipeline execution', 'Executes a complete CI/CD pipeline'),
        ('POST /track', 'Progress tracking', 'Tracks progress of ongoing operations')
    ]
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    header = table.rows[0].cells
    header[0].text = 'Endpoint'
    header[1].text = 'Purpose'
    header[2].text = 'Description'
    for cell in header:
        cell.paragraphs[0].runs[0].bold = True
    
    for endpoint, purpose, desc in endpoints:
        row = table.add_row().cells
        row[0].text = endpoint
        row[1].text = purpose
        row[2].text = desc
    
    # 5. Autonomous DevOps Agent Workflow
    doc.add_page_break()
    doc.add_heading('5. Autonomous DevOps Agent Workflow', level=1)
    
    doc.add_paragraph(
        'The autonomous agent is the flagship feature of this POC, demonstrating how AI can '
        'orchestrate complete DevOps workflows without human intervention.'
    )
    
    # Add workflow diagram
    diagram_path = 'generated-diagrams/autonomous-agent-workflow.png'
    if os.path.exists(diagram_path):
        doc.add_paragraph()
        doc.add_picture(diagram_path, width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption = doc.add_paragraph('Figure 3: Autonomous Agent Workflow')
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.runs[0].italic = True
    
    doc.add_paragraph()
    doc.add_heading('Workflow Steps:', level=2)
    
    steps = [
        ('Step 1: Repository Analysis', 
         'The agent clones the target repository and performs comprehensive analysis including:\n'
         '- Language detection and file structure mapping\n'
         '- Dependency analysis\n'
         '- Code pattern identification\n'
         '- AI-powered code review'),
        
        ('Step 2: Code Generation',
         'Based on user requirements or identified needs, the agent:\n'
         '- Generates new code following project conventions\n'
         '- Creates necessary configuration files\n'
         '- Adds appropriate documentation'),
        
        ('Step 3: Security Scanning',
         'The agent runs comprehensive security scans:\n'
         '- Trivy for container and dependency vulnerabilities\n'
         '- Bandit for Python security issues\n'
         '- Custom pattern matching for common security anti-patterns'),
        
        ('Step 4: Vulnerability Fixing',
         'For each identified vulnerability:\n'
         '- AI analyzes the root cause\n'
         '- Generates appropriate fix\n'
         '- Validates the fix doesn\'t break existing functionality'),
        
        ('Step 5: Pull Request Creation',
         'The agent creates a pull request with:\n'
         '- Detailed description of changes\n'
         '- List of fixed vulnerabilities\n'
         '- Before/after comparisons\n'
         '- Automated labels and reviewers'),
        
        ('Step 6: Tracking & Monitoring',
         'The agent provides:\n'
         '- Real-time progress updates\n'
         '- Detailed logs of all operations\n'
         '- Summary reports of completed actions')
    ]
    
    for step_title, step_desc in steps:
        doc.add_heading(step_title, level=3)
        doc.add_paragraph(step_desc)
    
    # 6. CI/CD Pipeline Integration
    doc.add_page_break()
    doc.add_heading('6. CI/CD Pipeline Integration', level=1)
    
    doc.add_paragraph(
        'The system integrates seamlessly with GitHub Actions to provide automated CI/CD pipelines.'
    )
    
    # Add CI/CD diagram
    diagram_path = 'generated-diagrams/cicd-pipeline-flow.png'
    if os.path.exists(diagram_path):
        doc.add_paragraph()
        doc.add_picture(diagram_path, width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption = doc.add_paragraph('Figure 4: CI/CD Pipeline Flow')
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.runs[0].italic = True
    
    doc.add_paragraph()
    doc.add_heading('Pipeline Stages:', level=2)
    
    pipeline_stages = [
        ('Build', 'Compiles code and creates container images'),
        ('Test', 'Runs unit tests, integration tests, and code quality checks'),
        ('Security Scan', 'Performs Trivy and Bandit security scans'),
        ('AI Analysis', 'Uses Bedrock to analyze code changes and provide recommendations'),
        ('Deploy', 'Deploys to AWS infrastructure using Lambda and API Gateway')
    ]
    
    for stage, desc in pipeline_stages:
        p = doc.add_paragraph()
        p.add_run(f'{stage}: ').bold = True
        p.add_run(desc)
    
    # 7. Security Features
    doc.add_page_break()
    doc.add_heading('7. Security Features', level=1)
    
    doc.add_heading('7.1 Trivy Integration', level=2)
    doc.add_paragraph(
        'Trivy is a comprehensive vulnerability scanner that checks for:'
    )
    security_items = [
        'Container image vulnerabilities',
        'Filesystem vulnerabilities',
        'Git repository issues',
        'Kubernetes misconfigurations',
        'Infrastructure as Code (IaC) issues'
    ]
    for item in security_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('7.2 Bandit Integration', level=2)
    doc.add_paragraph(
        'Bandit is a Python security linter that identifies:'
    )
    bandit_items = [
        'Hardcoded passwords and secrets',
        'SQL injection vulnerabilities',
        'Command injection risks',
        'Insecure cryptographic practices',
        'Dangerous function usage'
    ]
    for item in bandit_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('7.3 AI-Powered Security Analysis', level=2)
    doc.add_paragraph(
        'Beyond traditional scanning, the AI provides:'
    )
    ai_security = [
        'Context-aware vulnerability assessment',
        'Risk prioritization based on business impact',
        'Intelligent fix recommendations',
        'Security best practice suggestions'
    ]
    for item in ai_security:
        doc.add_paragraph(item, style='List Bullet')
    
    # 8. Technology Stack
    doc.add_page_break()
    doc.add_heading('8. Technology Stack', level=1)
    
    tech_categories = [
        ('Cloud Platform', ['AWS Lambda', 'AWS API Gateway', 'Amazon Bedrock', 'AWS CloudWatch']),
        ('AI/ML', ['Amazon Nova Pro (v1:0)', 'Natural Language Processing', 'Code Generation Models']),
        ('Programming Languages', ['Python 3.12', 'JavaScript (ES6+)', 'HTML5/CSS3']),
        ('Security Tools', ['Trivy', 'Bandit', 'Custom Security Scanners']),
        ('Version Control', ['Git', 'GitHub API', 'GitHub Actions']),
        ('Frontend', ['Vanilla JavaScript', 'CSS Grid/Flexbox', 'Fetch API'])
    ]
    
    for category, items in tech_categories:
        doc.add_heading(category, level=2)
        for item in items:
            doc.add_paragraph(item, style='List Bullet')
    
    # 9. Deployment Guide
    doc.add_page_break()
    doc.add_heading('9. Deployment Guide', level=1)
    
    doc.add_heading('9.1 Prerequisites', level=2)
    prereqs = [
        'AWS Account with appropriate permissions',
        'GitHub account and personal access token',
        'Python 3.12 or later',
        'AWS CLI configured with credentials'
    ]
    for prereq in prereqs:
        doc.add_paragraph(prereq, style='List Bullet')
    
    doc.add_heading('9.2 Lambda Deployment', level=2)
    doc.add_paragraph('Deploy the Lambda function with the following configuration:')
    
    config_items = [
        ('Runtime', 'Python 3.12'),
        ('Memory', '256 MB'),
        ('Timeout', '300 seconds (5 minutes)'),
        ('Handler', 'index.handler'),
        ('Required Files', 'index.py, opensource_tools.py')
    ]
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    header = table.rows[0].cells
    header[0].text = 'Setting'
    header[1].text = 'Value'
    for cell in header:
        cell.paragraphs[0].runs[0].bold = True
    
    for setting, value in config_items:
        row = table.add_row().cells
        row[0].text = setting
        row[1].text = value
    
    doc.add_paragraph()
    
    doc.add_heading('9.3 Environment Variables', level=2)
    env_vars = [
        ('GITHUB_TOKEN', 'GitHub personal access token'),
        ('AWS_REGION', 'AWS region (e.g., us-east-1)'),
        ('BEDROCK_MODEL_ID', 'amazon.nova-pro-v1:0')
    ]
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    header = table.rows[0].cells
    header[0].text = 'Variable'
    header[1].text = 'Description'
    for cell in header:
        cell.paragraphs[0].runs[0].bold = True
    
    for var, desc in env_vars:
        row = table.add_row().cells
        row[0].text = var
        row[1].text = desc
    
    # 10. Usage Examples
    doc.add_page_break()
    doc.add_heading('10. Usage Examples', level=1)
    
    doc.add_heading('10.1 Analyzing a Repository', level=2)
    doc.add_paragraph('Request:')
    doc.add_paragraph(
        'POST /analyze\n'
        '{\n'
        '  "repo_url": "https://github.com/owner/repo",\n'
        '  "branch": "main"\n'
        '}',
        style='No Spacing'
    )
    
    doc.add_heading('10.2 Generating Code', level=2)
    doc.add_paragraph('Request:')
    doc.add_paragraph(
        'POST /generate\n'
        '{\n'
        '  "requirement": "Create a Python function to validate email addresses",\n'
        '  "language": "python"\n'
        '}',
        style='No Spacing'
    )
    
    doc.add_heading('10.3 Running Security Scan', level=2)
    doc.add_paragraph('Request:')
    doc.add_paragraph(
        'POST /scan\n'
        '{\n'
        '  "repo_url": "https://github.com/owner/repo",\n'
        '  "scan_type": "full"\n'
        '}',
        style='No Spacing'
    )
    
    doc.add_heading('10.4 Running Autonomous Agent', level=2)
    doc.add_paragraph('Request:')
    doc.add_paragraph(
        'POST /autonomous\n'
        '{\n'
        '  "repo_url": "https://github.com/owner/repo",\n'
        '  "task": "Analyze repository, fix security issues, and create PR",\n'
        '  "create_pr": true\n'
        '}',
        style='No Spacing'
    )
    
    # Conclusion
    doc.add_page_break()
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph(
        'The BCG DevOps GenAI POC demonstrates the transformative potential of integrating '
        'Generative AI into DevOps workflows. By automating code analysis, security scanning, '
        'vulnerability fixing, and pull request creation, this solution significantly reduces '
        'manual effort while improving code quality and security posture.'
    )
    doc.add_paragraph()
    doc.add_paragraph(
        'Key benefits include:'
    )
    benefits = [
        'Reduced time from code commit to production deployment',
        'Consistent security practices across all code changes',
        'AI-powered insights for better code quality',
        'Automated documentation and change tracking',
        'Seamless integration with existing GitHub workflows'
    ]
    for benefit in benefits:
        doc.add_paragraph(benefit, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph(
        'This POC serves as a foundation for building more sophisticated AI-powered DevOps '
        'solutions that can adapt to specific organizational needs and workflows.'
    )
    
    # Save document
    output_path = 'BCG_DevOps_GenAI_POC_Documentation.docx'
    doc.save(output_path)
    print(f'Documentation saved to: {output_path}')
    return output_path

if __name__ == '__main__':
    create_documentation()
