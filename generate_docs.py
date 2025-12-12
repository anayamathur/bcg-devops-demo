#!/usr/bin/env python3
"""
Generate professional Word document for BCG DevOps + GenAI PoC
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime

def create_bcg_document():
    """Create professional Word document"""
    doc = Document()
    
    # Set document styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title
    title = doc.add_heading('BCG DevOps + GenAI Agentic Solution', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph('Proof of Concept Documentation')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    subtitle.runs[0].font.size = Pt(14)
    
    # Meta info
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f'Version 2.0 | December 2024\n').italic = True
    meta.add_run('Built on AWS Bedrock with Nova Pro AI').italic = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Table of Contents
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Executive Summary',
        '2. Architecture Overview',
        '3. Key Features',
        '4. API Reference',
        '5. Autonomous Agent',
        '6. Quick Start Guide',
        '7. Integration Guide',
        '8. Security Policies',
        '9. Sample Conversations',
        '10. Cost Estimation',
        '11. Troubleshooting',
        '12. Support'
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_page_break()
    
    # Executive Summary
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        'This Proof of Concept demonstrates an Agentic AI-powered DevOps platform built on '
        'AWS Bedrock that revolutionizes how development teams interact with CI/CD pipelines. '
        'The solution enables natural language conversations to manage, create, and maintain '
        'DevOps workflows.'
    )
    
    # Capabilities table
    doc.add_heading('Key Capabilities', level=2)
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Table Grid'
    
    capabilities = [
        ('Capability', 'Description'),
        ('Conversational DevOps', 'Chat-based interface for all DevOps operations'),
        ('Intelligent Template Generation', 'Auto-generate CI/CD workflows for any tech stack'),
        ('Autonomous Agent', 'Self-healing CI/CD - auto-fix failures, retry, and deploy'),
        ('Security-First Approach', 'AI-powered security gatekeeper for all changes'),
        ('Full Repository Awareness', 'Agents understand complete codebase context'),
        ('Incident Response', 'L1 auto-remediation, L2/L3 human-in-loop'),
    ]
    
    for i, (cap, desc) in enumerate(capabilities):
        row = table.rows[i]
        row.cells[0].text = cap
        row.cells[1].text = desc
        if i == 0:
            for cell in row.cells:
                cell.paragraphs[0].runs[0].bold = True
    
    doc.add_page_break()
    
    # Architecture Overview
    doc.add_heading('2. Architecture Overview', level=1)
    
    doc.add_heading('System Architecture', level=2)
    doc.add_paragraph(
        'The platform consists of several AWS services working together to provide '
        'an intelligent DevOps assistant:'
    )
    
    arch_text = """
    User Interface (Web/Slack)
           |
           v
    API Gateway (REST API)
           |
           v
    Lambda Function (Python 3.12)
           |
    +------+------+------+------+
    |      |      |      |      |
    v      v      v      v      v
 Bedrock  GitHub  Secrets  CloudWatch
 Nova Pro  API    Manager   Logs
    """
    
    arch_para = doc.add_paragraph()
    arch_run = arch_para.add_run(arch_text)
    arch_run.font.name = 'Courier New'
    arch_run.font.size = Pt(9)
    
    doc.add_heading('Core Components', level=2)
    comp_table = doc.add_table(rows=7, cols=3)
    comp_table.style = 'Table Grid'
    
    components = [
        ('Component', 'AWS Service', 'Purpose'),
        ('DevOps Agent', 'Lambda + Bedrock', 'Main orchestrator for CI/CD operations'),
        ('AI Engine', 'Bedrock Nova Pro', 'Natural language processing and generation'),
        ('API Layer', 'API Gateway', 'REST endpoints for all operations'),
        ('Repository Integration', 'GitHub API', 'Clone, commit, PR, Actions tracking'),
        ('Secrets Management', 'Secrets Manager', 'Secure credential storage'),
        ('Monitoring', 'CloudWatch', 'Logs and metrics'),
    ]
    
    for i, (comp, service, purpose) in enumerate(components):
        row = comp_table.rows[i]
        row.cells[0].text = comp
        row.cells[1].text = service
        row.cells[2].text = purpose
        if i == 0:
            for cell in row.cells:
                cell.paragraphs[0].runs[0].bold = True
    
    doc.add_page_break()
    
    # Key Features
    doc.add_heading('3. Key Features', level=1)
    
    doc.add_heading('3.1 Autonomous DevOps Agent', level=2)
    doc.add_paragraph(
        'The crown jewel of the solution - a fully autonomous agent that can:'
    )
    features = [
        'Analyze repositories and detect tech stack automatically',
        'Understand natural language requests',
        'Generate optimized CI/CD workflows',
        'Create branches and commit changes',
        'Open Pull Requests with proper descriptions',
        'Monitor GitHub Actions execution',
        'Auto-fix failures and retry operations',
        'Report final status with detailed logs'
    ]
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_heading('3.2 Dynamic Chat Agent', level=2)
    doc.add_paragraph(
        'Developers can interact with the platform using natural language. '
        'The agent understands context, analyzes repositories, and provides '
        'intelligent responses.'
    )
    
    doc.add_heading('3.3 Intelligent Template Generation', level=2)
    gen_features = [
        'Auto-detects: Language, framework, package manager',
        'Generates: Optimized CI/CD workflows with caching',
        'Validates: Security best practices enforcement',
        'Integrates: BCG-approved tools (SonarQube, JFrog, Prisma)'
    ]
    for f in gen_features:
        doc.add_paragraph(f, style='List Bullet')
    
    doc.add_heading('3.4 Self-Healing CI/CD', level=2)
    doc.add_paragraph('When workflows fail, the agent automatically:')
    healing_steps = [
        'Fetches failure logs from GitHub Actions',
        'Identifies failed steps and jobs',
        'Generates intelligent fixes using AI',
        'Commits fix and retries automatically',
        'Reports final status'
    ]
    for i, step in enumerate(healing_steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_page_break()
    
    # API Reference
    doc.add_heading('4. API Reference', level=1)
    
    doc.add_heading('Base URL', level=2)
    doc.add_paragraph('https://4dyb4z9kgk.execute-api.us-east-1.amazonaws.com/prod')
    
    doc.add_heading('Endpoints', level=2)
    
    endpoints = [
        ('GET /health', 'Health check - verify API status'),
        ('POST /analyze', 'Analyze repository and detect tech stack'),
        ('POST /generate', 'Generate workflow for repository'),
        ('POST /create-pr', 'Create PR with workflow file'),
        ('POST /track', 'Track GitHub Actions status'),
        ('POST /suggest', 'Get intelligent suggestions'),
        ('POST /fix', 'Auto-fix failing workflows'),
        ('POST /validate', 'Validate workflow content'),
        ('POST /knowledge', 'Build comprehensive project knowledge'),
        ('POST /autonomous', 'Run fully autonomous agent'),
    ]
    
    endpoint_table = doc.add_table(rows=len(endpoints)+1, cols=2)
    endpoint_table.style = 'Table Grid'
    
    header = endpoint_table.rows[0]
    header.cells[0].text = 'Endpoint'
    header.cells[1].text = 'Description'
    for cell in header.cells:
        cell.paragraphs[0].runs[0].bold = True
    
    for i, (endpoint, desc) in enumerate(endpoints, 1):
        row = endpoint_table.rows[i]
        row.cells[0].text = endpoint
        row.cells[1].text = desc
    
    doc.add_page_break()
    
    # Autonomous Agent
    doc.add_heading('5. Autonomous Agent', level=1)
    
    doc.add_heading('The /autonomous Endpoint', level=2)
    doc.add_paragraph(
        'This is the most powerful endpoint - a fully autonomous DevOps agent that can '
        'handle complex requests end-to-end without human intervention.'
    )
    
    doc.add_heading('Request Format', level=3)
    request_example = '''POST /autonomous
Content-Type: application/json

{
  "repository": "owner/repo",
  "request": "Add CI/CD pipeline with testing and deployment",
  "max_retries": 3,
  "auto_merge": false
}'''
    req_para = doc.add_paragraph()
    req_run = req_para.add_run(request_example)
    req_run.font.name = 'Courier New'
    req_run.font.size = Pt(9)
    
    doc.add_heading('Execution Flow', level=3)
    flow_steps = [
        'ANALYZE - Scans repository, detects tech stack',
        'UNDERSTAND - AI processes request, plans action',
        'GENERATE - Creates optimized workflow YAML',
        'DEPLOY - Creates branch, commits, opens PR',
        'MONITOR - Watches GitHub Actions run',
        'EVALUATE - Determines success or failure',
        'AUTO-FIX - If failed, generates fix',
        'RETRY - Up to max_retries attempts'
    ]
    for i, step in enumerate(flow_steps, 1):
        doc.add_paragraph(f'Step {i}: {step}')
    
    doc.add_page_break()
    
    # Quick Start
    doc.add_heading('6. Quick Start Guide', level=1)
    
    doc.add_heading('Prerequisites', level=2)
    prereqs = [
        'AWS Account with Bedrock access (Nova Pro enabled)',
        'Terraform >= 1.5.0',
        'AWS CLI configured',
        'GitHub Personal Access Token'
    ]
    for p in prereqs:
        doc.add_paragraph(p, style='List Bullet')
    
    doc.add_heading('Deployment Steps', level=2)
    
    deploy_code = '''# Clone repository
git clone https://github.com/your-org/bcg-devops-genai-poc.git
cd bcg-devops-genai-poc

# Deploy infrastructure
cd infrastructure/terraform
terraform init
terraform plan -var="aws_region=us-east-1"
terraform apply'''
    
    deploy_para = doc.add_paragraph()
    deploy_run = deploy_para.add_run(deploy_code)
    deploy_run.font.name = 'Courier New'
    deploy_run.font.size = Pt(9)
    
    doc.add_heading('Configure GitHub Token', level=2)
    token_code = '''aws secretsmanager put-secret-value \\
  --secret-id bcg-devops-genai/github-token \\
  --secret-string '{"token": "ghp_YOUR_TOKEN_HERE"}' \\
  --region us-east-1'''
    
    token_para = doc.add_paragraph()
    token_run = token_para.add_run(token_code)
    token_run.font.name = 'Courier New'
    token_run.font.size = Pt(9)
    
    doc.add_page_break()
    
    # Integration Guide
    doc.add_heading('7. Integration Guide', level=1)
    
    doc.add_heading('BCG Standard Tools Integration', level=2)
    
    tools_table = doc.add_table(rows=8, cols=3)
    tools_table.style = 'Table Grid'
    
    tools = [
        ('Tool', 'Integration Method', 'Purpose'),
        ('SonarQube', 'sonarqube-scan-action', 'Code quality gates'),
        ('JFrog Artifactory', 'jfrog/setup-jfrog-cli', 'Artifact storage'),
        ('Prisma Cloud', 'prisma-cloud-scan', 'Security scanning'),
        ('Trivy', 'aquasecurity/trivy-action', 'Container scanning'),
        ('ArgoCD', 'argocd-sync', 'GitOps deployment'),
        ('Datadog', 'datadog-agent', 'Monitoring'),
        ('Slack', 'slack-notify', 'Notifications'),
    ]
    
    for i, (tool, method, purpose) in enumerate(tools):
        row = tools_table.rows[i]
        row.cells[0].text = tool
        row.cells[1].text = method
        row.cells[2].text = purpose
        if i == 0:
            for cell in row.cells:
                cell.paragraphs[0].runs[0].bold = True
    
    doc.add_page_break()
    
    # Security Policies
    doc.add_heading('8. Security Policies', level=1)
    
    doc.add_heading('Workflow Security Requirements', level=2)
    
    security_table = doc.add_table(rows=6, cols=2)
    security_table.style = 'Table Grid'
    
    security_items = [
        ('Category', 'Requirement'),
        ('Secrets', 'Use ${{ secrets.NAME }} - no hardcoded values'),
        ('Action Versions', 'Pin to specific versions (@v4, not @main)'),
        ('Permissions', 'Use least privilege principle'),
        ('Scanning', 'Required for all workflows'),
        ('Production Deploy', 'Requires manual approval'),
    ]
    
    for i, (cat, req) in enumerate(security_items):
        row = security_table.rows[i]
        row.cells[0].text = cat
        row.cells[1].text = req
        if i == 0:
            for cell in row.cells:
                cell.paragraphs[0].runs[0].bold = True
    
    doc.add_heading('Required Security Checks', level=2)
    checks = [
        'All tests pass',
        'Code coverage >= 80%',
        'Security scan passes (no HIGH/CRITICAL)',
        'SonarQube quality gate passes',
        'Peer review approved'
    ]
    for check in checks:
        doc.add_paragraph(f'[ ] {check}')
    
    doc.add_page_break()
    
    # Sample Conversations
    doc.add_heading('9. Sample Conversations', level=1)
    
    doc.add_heading('Example 1: New Project Setup', level=2)
    
    convo1 = '''User: I need a CI/CD pipeline for my Python FastAPI project at 
      github.com/bcg/user-service. Include tests, security scanning,
      Docker build, and EKS deployment.

Agent: Analyzing github.com/bcg/user-service...

       Detected:
       - Python 3.11 with FastAPI
       - Poetry for dependencies
       - Dockerfile present (multi-stage)
       - Kubernetes manifests in /k8s
       
       Creating workflow with:
       1. Lint & Format (ruff, black)
       2. Unit Tests (pytest with coverage)
       3. Security Scan (Trivy + SonarQube)
       4. Build Docker image
       5. Push to JFrog Artifactory
       6. Deploy to EKS via ArgoCD
       
       PR Created: https://github.com/bcg/user-service/pull/42'''
    
    convo1_para = doc.add_paragraph()
    convo1_run = convo1_para.add_run(convo1)
    convo1_run.font.name = 'Courier New'
    convo1_run.font.size = Pt(9)
    
    doc.add_heading('Example 2: Fix Failing Pipeline', level=2)
    
    convo2 = '''User: My CI is failing, can you fix it?

Agent: Analyzing recent workflow runs...

       Found Failure:
       - Run ID: #156
       - Failed Step: "Install dependencies"
       - Error: "npm ERR! ERESOLVE could not resolve"
       
       Root Cause:
       - Dependency conflict in package-lock.json
       
       Fix Applied:
       - Updated npm ci to npm ci --legacy-peer-deps
       - Added Node.js version matrix
       
       PR Created: https://github.com/bcg/app/pull/157
       
       Monitoring new run... SUCCESS!'''
    
    convo2_para = doc.add_paragraph()
    convo2_run = convo2_para.add_run(convo2)
    convo2_run.font.name = 'Courier New'
    convo2_run.font.size = Pt(9)
    
    doc.add_page_break()
    
    # Cost Estimation
    doc.add_heading('10. Cost Estimation', level=1)
    
    doc.add_heading('Monthly Estimate (Production)', level=2)
    
    cost_table = doc.add_table(rows=7, cols=3)
    cost_table.style = 'Table Grid'
    
    costs = [
        ('Service', 'Usage', 'Estimated Cost'),
        ('Bedrock (Nova Pro)', '1M input + 500K output tokens', '$18.00'),
        ('Lambda', '1M invocations, 512MB', '$15.00'),
        ('API Gateway', '1M requests', '$3.50'),
        ('Secrets Manager', '5 secrets', '$2.00'),
        ('CloudWatch', 'Logs + Metrics', '$30.00'),
        ('Total', '', '~$70/month'),
    ]
    
    for i, (service, usage, cost) in enumerate(costs):
        row = cost_table.rows[i]
        row.cells[0].text = service
        row.cells[1].text = usage
        row.cells[2].text = cost
        if i == 0 or i == 6:
            for cell in row.cells:
                cell.paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    doc.add_paragraph(
        'Note: Costs scale with usage. Enterprise deployments with OpenSearch may be higher.'
    ).italic = True
    
    doc.add_page_break()
    
    # Troubleshooting
    doc.add_heading('11. Troubleshooting', level=1)
    
    doc.add_heading('Common Issues', level=2)
    
    doc.add_heading('1. GitHub Token Not Working', level=3)
    doc.add_paragraph('Verify token in Secrets Manager and ensure it has these scopes:')
    doc.add_paragraph('repo (full control)', style='List Bullet')
    doc.add_paragraph('workflow (Actions access)', style='List Bullet')
    
    doc.add_heading('2. Bedrock Model Access', level=3)
    doc.add_paragraph(
        'Ensure Nova Pro is enabled in us-east-1 region. Check model access in Bedrock console.'
    )
    
    doc.add_heading('3. Lambda Timeout', level=3)
    doc.add_paragraph(
        'For autonomous operations, increase Lambda timeout to 300 seconds.'
    )
    
    doc.add_heading('4. Workflow Not Triggering', level=3)
    doc.add_paragraph('Check that .github/workflows/ directory exists', style='List Bullet')
    doc.add_paragraph('Verify workflow YAML syntax', style='List Bullet')
    doc.add_paragraph('Ensure GitHub Actions is enabled for repository', style='List Bullet')
    
    doc.add_page_break()
    
    # Support
    doc.add_heading('12. Support', level=1)
    
    doc.add_paragraph('For questions or issues, contact the i2k2 team:')
    
    support_table = doc.add_table(rows=4, cols=2)
    support_table.style = 'Table Grid'
    
    support_info = [
        ('Role', 'Contact'),
        ('Architecture', 'i2k2 Networks'),
        ('Implementation', 'DevOps Team'),
        ('AWS Support', 'AWS SA'),
    ]
    
    for i, (role, contact) in enumerate(support_info):
        row = support_table.rows[i]
        row.cells[0].text = role
        row.cells[1].text = contact
        if i == 0:
            for cell in row.cells:
                cell.paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Footer
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run('Last Updated: December 2024').italic = True
    footer.add_run('\nBuilt with AWS Bedrock Nova Pro').italic = True
    
    # Save document
    doc_path = '/home/raj/Desktop/allwork/bcg-devops-genai-poc/docs/BCG_DevOps_GenAI_PoC_Documentation.docx'
    doc.save(doc_path)
    print(f'Document saved to: {doc_path}')
    return doc_path

if __name__ == '__main__':
    create_bcg_document()
